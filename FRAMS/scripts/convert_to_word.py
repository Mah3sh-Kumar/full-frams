
import os
import glob
from docx import Document
from docx.shared import Pt
import re

def convert_md_to_docx(source_dir, output_file):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Get all markdown files in order
    files = sorted(glob.glob(os.path.join(source_dir, "Chapter_*.md")), 
                   key=lambda x: int(re.search(r'Chapter_(\d+)', os.path.basename(x)).group(1)))
    
    if not files:
        print("No Chapter files found in", source_dir)
        return

    print("Found files:", [os.path.basename(f) for f in files])

    for i, file_path in enumerate(files):
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Process lines
        code_block = False
        table_mode = False
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines
            if not line:
                continue
                
            # Headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
                
            # Lists
            elif line.startswith('* ') or line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                # Remove the number and dot (e.g., "1. " -> "")
                text = re.sub(r'^\d+\. ', '', line)
                p = doc.add_paragraph(text, style='List Number')
                
            # Code Blocks
            elif line.startswith('```'):
                code_block = not code_block
                continue
            elif code_block:
                p = doc.add_paragraph(line)
                p.style = 'No Spacing'
                p.runs[0].font.name = 'Courier New'
                
            # Tables (Basic support - just printing raw pipes for now as converting md tables to docx tables is complex)
            elif line.startswith('|'):
                p = doc.add_paragraph(line)
                p.runs[0].font.name = 'Courier New'
            
            # Regular Paragraph
            else:
                doc.add_paragraph(line)
        
        # Add Page Break after each chapter, except the last one
        if i < len(files) - 1:
            doc.add_page_break()

    doc.save(output_file)
    print(f"Successfully created {output_file}")

if __name__ == "__main__":
    source_directory = r"e:\\SEM_6_project\\FRAMS\\DOCUMENTATIONS"
    output_filename = r"e:\\SEM_6_project\\FRAMS\\DOCUMENTATIONS\\FRAMS_Final_Report.docx"
    
    convert_md_to_docx(source_directory, output_filename)
